import csv
import logging
import os
import shutil
import tempfile
import zipfile
from typing import Dict, List, Tuple

from django.core.files.storage import default_storage
from django.http import HttpResponse
from django.utils.encoding import smart_str

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl

from .models import Document
import logging

docx_logger = logging.getLogger("docx_export")


def export_to_txt(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в TXT файл
    """
    sentences = document.sentences.all().order_by("sentence_number")

    with open(output_path, "w", encoding="utf-8") as file:
        for sentence in sentences:
            original_text = sentence.original_text or ""
            translated_text = sentence.translation.translated_text if sentence.has_translation else ""
            file.write(f"{original_text}\t{translated_text}\n")

    return output_path


def _create_tc(text: str) -> OxmlElement:
    tc = OxmlElement("w:tc")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    text = text or ""
    # Сохраняем пробелы, если есть ведущие/замыкающие
    if text.strip() != text:
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _append_row_fast(table, values):
    tr = OxmlElement("w:tr")
    for v in values:
        tr.append(_create_tc(str(v) if v is not None else ""))
    table._tbl.append(tr)


def export_to_docx(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в DOCX файл в виде таблицы
    """
    docx_logger.info(f"Начало экспорта DOCX (таблица) для документа id={document.id} title='{document.title}'")
    doc = docx.Document()

    sentences = list(document.sentences.all().order_by("sentence_number"))
    total_rows = len(sentences)
    docx_logger.info(f"Создание таблицы на {total_rows} строк (+1 заголовок)")

    # Таблица с заголовками: №, Оригинал, Перевод
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "№"
    header_cells[1].text = "Оригинал"
    header_cells[2].text = "Перевод"

    # Быстрый режим добавления строк через oxml (значительно быстрее для больших таблиц)
    for idx, sentence in enumerate(sentences, start=1):
        orig = sentence.original_text or ""
        trans = sentence.translation.translated_text if sentence.has_translation else ""
        _append_row_fast(table, [sentence.sentence_number, orig, trans])
        if idx % 1000 == 0:
            docx_logger.info(f"Заполнено строк: {idx}/{total_rows}")
    docx_logger.info(f"Сформирована таблица: строк={total_rows}; путь сохранения='{output_path}'")

    try:
        doc.save(output_path)
        docx_logger.info(f"DOCX (таблица) сохранен: '{output_path}'")
    except Exception as e:
        docx_logger.exception(f"Ошибка сохранения DOCX (таблица): {e}")
        raise
    return output_path


def _iter_all_paragraphs(document_obj) -> List[docx.text.paragraph.Paragraph]:
    paragraphs: List[docx.text.paragraph.Paragraph] = list(document_obj.paragraphs)
    # Обходим таблицы рекурсивно
    for table in document_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _replace_text_in_runs(paragraph: "docx.text.paragraph.Paragraph", replacements: List[Tuple[str, str]]):
    # Заменяем внутри каждого run, чтобы максимально сохранить форматирование run'ов
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        for src, dst in replacements:
            if src and src in text:
                text = text.replace(src, dst)
        run.text = text


def export_to_docx_translated_only(document: Document, original_docx_path: str, output_path: str) -> str:
    """
    Создает DOCX, повторяющий структуру исходного документа, заменяя исходные предложения переведенным текстом.
    Примечание: сохранение форматирования ограничено — при замене внутри run форматирование сохраняется,
    но предложения, разбитые на несколько run, могут быть заменены частично.
    """
    # Словарь замен: оригинал -> перевод
    sentence_map: List[Tuple[str, str]] = []
    for s in document.sentences.all().order_by("sentence_number"):
        if hasattr(s, "translation") and s.translation and s.translation.translated_text:
            # более длинные исходные строки должны заменяться сначала
            sentence_map.append((s.original_text, s.translation.translated_text))
    # Сортируем по длине исходного текста по убыванию
    sentence_map.sort(key=lambda x: len(x[0] or ""), reverse=True)

    doc = docx.Document(original_docx_path)

    # Параграфы документа и ячейки таблиц
    for paragraph in _iter_all_paragraphs(doc):
        _replace_text_in_runs(paragraph, sentence_map)

    doc.save(output_path)
    return output_path


def export_to_xlsx(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в XLSX файл
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Переводы"

    # Заголовки
    headers = [
        "№",
        "Оригинальный текст",
        "Переведенный текст",
        "Переводчик",
        "Корректор",
        "Статус",
        "Дата перевода",
        "Дата корректировки",
    ]

    for col, header in enumerate(headers, 1):
        sheet.cell(row=1, column=col, value=header)

    # Данные
    sentences = document.sentences.all().order_by("sentence_number")

    for row, sentence in enumerate(sentences, 2):
        sheet.cell(row=row, column=1, value=sentence.sentence_number)
        sheet.cell(row=row, column=2, value=sentence.original_text)

        if sentence.has_translation:
            translation = sentence.translation
            sheet.cell(row=row, column=3, value=translation.translated_text)
            sheet.cell(row=row, column=4, value=str(translation.translator))
            sheet.cell(
                row=row,
                column=5,
                value=str(sentence.corrector) if sentence.corrector else "",
            )

            # Статус
            if translation.status == "approved":
                status = "Одобрен"
            elif translation.status == "rejected":
                status = "Отклонен"
            else:
                status = "На проверке"
            sheet.cell(row=row, column=6, value=status)
            sheet.cell(
                row=row,
                column=7,
                value=translation.translated_at.strftime("%Y-%m-%d %H:%M"),
            )
            sheet.cell(
                row=row,
                column=8,
                value=(translation.corrected_at.strftime("%Y-%m-%d %H:%M") if translation.corrected_at else ""),
            )
        else:
            sheet.cell(row=row, column=3, value="Не переведено")

    # Автоматическая ширина столбцов
    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except Exception:
                continue
        adjusted_width = min(max_length + 2, 50)
        sheet.column_dimensions[column_letter].width = adjusted_width

    workbook.save(output_path)
    return output_path


def export_document_translations(document: Document, format_type: str) -> str:
    """
    Основная функция для экспорта переводов документа
    """
    # Создаем папку для экспорта, если её нет
    export_dir = os.path.join(default_storage.location, "exports")
    os.makedirs(export_dir, exist_ok=True)

    # Генерируем имя файла
    base_name = f"translation_{document.id}_{document.title.replace(' ', '_')}"

    if format_type == "txt":
        output_path = os.path.join(export_dir, f"{base_name}.txt")
        return export_to_txt(document, output_path)
    elif format_type == "docx_table":
        output_path = os.path.join(export_dir, f"{base_name}_table.docx")
        return export_to_docx(document, output_path)
    elif format_type == "docx_translated":
        output_path = os.path.join(export_dir, f"{base_name}_translated_only.docx")
        # Пытаемся использовать исходный DOCX для максимально точного воспроизведения
        original_path = None
        try:
            if document.file and document.file.name.lower().endswith(".docx"):
                original_path = default_storage.path(document.file.name)
        except Exception:
            original_path = None

        if original_path and os.path.exists(original_path):
            return export_to_docx_translated_only(document, original_path, output_path)
        else:
            # Фолбэк: создаем простой документ с переведенным текстом построчно
            simple_doc = docx.Document()
            for s in document.sentences.all().order_by("sentence_number"):
                text = s.translation.translated_text if hasattr(s, "translation") and s.translation else ""
                simple_doc.add_paragraph(text)
            simple_doc.save(output_path)
            return output_path
    elif format_type == "xlsx":
        output_path = os.path.join(export_dir, f"{base_name}.xlsx")
        return export_to_xlsx(document, output_path)
    else:
        raise ValueError(f"Неподдерживаемый формат экспорта: {format_type}")


def export_document_all_formats(document: Document) -> str:
    """
    Экспортирует документ во всех поддерживаемых форматах и упаковывает в ZIP архив
    """
    # Создаем временную директорию для файлов
    with tempfile.TemporaryDirectory() as temp_dir:
        # Генерируем базовое имя файла
        base_name = f"translation_{document.id}_{document.title.replace(' ', '_')}"

        # Экспортируем во все форматы
        txt_path = os.path.join(temp_dir, f"{base_name}.txt")
        docx_table_path = os.path.join(temp_dir, f"{base_name}_table.docx")
        docx_translated_path = os.path.join(temp_dir, f"{base_name}_translated_only.docx")
        xlsx_path = os.path.join(temp_dir, f"{base_name}.xlsx")

        # Создаем файлы
        export_to_txt(document, txt_path)
        export_to_docx(document, docx_table_path)
        # Сгенерировать translated_only DOCX
        original_path = None
        try:
            if document.file and document.file.name.lower().endswith(".docx"):
                original_path = default_storage.path(document.file.name)
        except Exception:
            original_path = None

        if original_path and os.path.exists(original_path):
            export_to_docx_translated_only(document, original_path, docx_translated_path)
        else:
            simple_doc = docx.Document()
            for s in document.sentences.all().order_by("sentence_number"):
                text = s.translation.translated_text if hasattr(s, "translation") and s.translation else ""
                simple_doc.add_paragraph(text)
            simple_doc.save(docx_translated_path)

        export_to_xlsx(document, xlsx_path)

        # Создаем ZIP архив
        zip_path = os.path.join(temp_dir, f"{base_name}_all_formats.zip")

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            # Добавляем файлы в архив
            zipf.write(txt_path, os.path.basename(txt_path))
            zipf.write(docx_table_path, os.path.basename(docx_table_path))
            zipf.write(docx_translated_path, os.path.basename(docx_translated_path))
            zipf.write(xlsx_path, os.path.basename(xlsx_path))

        # Копируем архив в постоянную директорию
        export_dir = os.path.join(default_storage.location, "exports")
        os.makedirs(export_dir, exist_ok=True)

        final_zip_path = os.path.join(export_dir, f"{base_name}_all_formats.zip")

        # Копируем файл
        shutil.copy2(zip_path, final_zip_path)

        return final_zip_path


def get_document_statistics(document: Document) -> Dict:
    """
    Возвращает статистику по документу
    """
    total_sentences = document.sentences.count()
    translated_sentences = document.sentences.filter(translation__isnull=False).count()
    approved_sentences = document.sentences.filter(translation__status="approved").count()
    rejected_sentences = document.sentences.filter(translation__status="rejected").count()
    pending_sentences = document.sentences.filter(translation__status="pending").count()

    # Статистика по словам и знакам в предложениях
    sentences = document.sentences.all()
    total_words = 0
    total_characters = 0
    total_characters_without_spaces = 0

    for sentence in sentences:
        # Подсчет слов (разделение по пробелам)
        words = sentence.original_text.split()
        total_words += len(words)

        # Подсчет символов с пробелами
        total_characters += len(sentence.original_text)

        # Подсчет символов без пробелов
        total_characters_without_spaces += len(sentence.original_text.replace(" ", ""))

    # Статистика по словам и знакам в переводах
    translations = document.sentences.filter(translation__isnull=False)
    total_translated_words = 0
    total_translated_characters = 0
    total_translated_characters_without_spaces = 0

    for sentence in translations:
        if sentence.translation:
            # Подсчет слов в переводе
            translated_words = sentence.translation.translated_text.split()
            total_translated_words += len(translated_words)

            # Подсчет символов в переводе
            total_translated_characters += len(sentence.translation.translated_text)

            # Подсчет символов без пробелов в переводе
            total_translated_characters_without_spaces += len(sentence.translation.translated_text.replace(" ", ""))

    # Средние показатели
    average_words_per_sentence = total_words / total_sentences if total_sentences > 0 else 0
    average_characters_per_sentence = total_characters / total_sentences if total_sentences > 0 else 0
    average_translated_words_per_translation = (
        total_translated_words / translated_sentences if translated_sentences > 0 else 0
    )
    average_translated_characters_per_translation = (
        total_translated_characters / translated_sentences if translated_sentences > 0 else 0
    )

    return {
        "total_sentences": total_sentences,
        "translated_sentences": translated_sentences,
        "approved_sentences": approved_sentences,
        "rejected_sentences": rejected_sentences,
        "pending_sentences": pending_sentences,
        "total_words": total_words,
        "total_characters": total_characters,
        "total_characters_without_spaces": total_characters_without_spaces,
        "total_translated_words": total_translated_words,
        "total_translated_characters": total_translated_characters,
        "total_translated_characters_without_spaces": total_translated_characters_without_spaces,
        "average_words_per_sentence": round(average_words_per_sentence, 1),
        "average_characters_per_sentence": round(average_characters_per_sentence, 1),
        "average_translated_words_per_translation": round(average_translated_words_per_translation, 1),
        "average_translated_characters_per_translation": round(average_translated_characters_per_translation, 1),
        "completion_percentage": round(
            ((translated_sentences / total_sentences * 100) if total_sentences > 0 else 0),
            2,
        ),
        "approval_percentage": round(
            (approved_sentences / total_sentences * 100) if total_sentences > 0 else 0,
            2,
        ),
    }


def export_sentences_to_csv(queryset):
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = "attachment; filename=sentences.csv"
    writer = csv.writer(response)
    writer.writerow(
        [
            "Документ",
            "№ предложения",
            "Оригинальный текст",
            "Статус",
            "Назначено",
            "Перевод",
            "Дата создания",
            "Дата обновления",
        ]
    )
    for s in queryset:
        writer.writerow(
            [
                smart_str(s.document.title),
                s.sentence_number,
                smart_str(s.original_text),
                s.get_status_display(),
                smart_str(s.assigned_to.get_full_name() if s.assigned_to else ""),
                smart_str(s.translation.translated_text if hasattr(s, "translation") else ""),
                s.created_at.strftime("%d.%m.%Y %H:%M"),
                s.updated_at.strftime("%d.%m.%Y %H:%M"),
            ]
        )
    return response


def export_user_report_to_xlsx(user_obj, user_stats, context_data, output_path: str) -> str:
    """
    Экспортирует отчет по пользователю в XLSX файл
    """
    try:
        workbook = openpyxl.Workbook()
        # Лист с основной информацией
        main_sheet = workbook.active
        main_sheet.title = "Основная информация"
        # Заголовок отчета
        main_sheet.merge_cells("A1:D1")
        title_cell = main_sheet["A1"]
        title_cell.value = f"Отчет по пользователю: {user_obj.get_full_name()}"
        title_cell.font = openpyxl.styles.Font(size=16, bold=True)
        title_cell.alignment = openpyxl.styles.Alignment(horizontal="center")
        # Основная информация о пользователе
        main_sheet["A3"] = "Основная информация"
        main_sheet["A3"].font = openpyxl.styles.Font(size=14, bold=True)
        user_info = [
            ("Полное имя", user_obj.get_full_name()),
            ("Имя пользователя", f"@{user_obj.username}"),
            ("Email", user_obj.email),
            ("Телефон", user_obj.phone or "Не указан"),
            ("Роль", user_obj.get_role_display()),
            ("Дата регистрации", user_obj.date_joined.strftime("%d.%m.%Y %H:%M")),
            (
                "Последний вход",
                (user_obj.last_login.strftime("%d.%m.%Y %H:%M") if user_obj.last_login else "Никогда"),
            ),
        ]
        for row, (label, value) in enumerate(user_info, 4):
            main_sheet[f"A{row}"] = label
            main_sheet[f"B{row}"] = value
            main_sheet[f"A{row}"].font = openpyxl.styles.Font(bold=True)
        # Статистика активности
        main_sheet["A12"] = "Статистика активности"
        main_sheet["A12"].font = openpyxl.styles.Font(size=14, bold=True)
        if user_obj.role in ["admin", "representative"]:
            activity_stats = [
                (
                    "Загруженных документов",
                    context_data.get("user_documents", []).count(),
                ),
            ]
        elif user_obj.role == "translator":
            activity_stats = [
                ("Всего назначено", context_data.get("total_assigned", 0)),
                ("В работе", context_data.get("in_progress_sentences", 0)),
                ("На проверке", context_data.get("total_translated", 0)),
                ("Завершено", context_data.get("total_completed", 0)),
                ("Отклонено", context_data.get("total_rejected", 0)),
            ]
        elif user_obj.role == "corrector":
            activity_stats = [
                ("Ожидают проверки", context_data.get("pending_corrections", 0)),
                ("Всего проверено", context_data.get("total_reviewed", 0)),
                ("Одобрено", context_data.get("approved_count", 0)),
                ("Отклонено", context_data.get("rejected_count", 0)),
            ]
        for row, (label, value) in enumerate(activity_stats, 13):
            main_sheet[f"A{row}"] = label
            main_sheet[f"B{row}"] = value
            main_sheet[f"A{row}"].font = openpyxl.styles.Font(bold=True)
        # Статистика переводов (только для переводчиков)
        if user_obj.role == "translator":
            main_sheet["A19"] = "Статистика переводов"
            main_sheet["A19"].font = openpyxl.styles.Font(size=14, bold=True)
            translation_stats = [
                ("Утверждено", context_data.get("approved_count", 0)),
                ("На проверке", context_data.get("pending_count", 0)),
                ("Отклонено", context_data.get("rejected_count", 0)),
            ]
            for row, (label, value) in enumerate(translation_stats, 20):
                main_sheet[f"A{row}"] = label
                main_sheet[f"B{row}"] = value
                main_sheet[f"A{row}"].font = openpyxl.styles.Font(bold=True)
        # Лист со статистикой по словам и символам
        stats_sheet = workbook.create_sheet("Статистика по словам и символам")
        # Заголовок
        stats_sheet.merge_cells("A1:E1")
        stats_title_cell = stats_sheet["A1"]
        stats_title_cell.value = "Статистика по словам и символам"
        stats_title_cell.font = openpyxl.styles.Font(size=16, bold=True)
        stats_title_cell.alignment = openpyxl.styles.Alignment(horizontal="center")
        # Заголовки таблицы
        headers = [
            "Тип",
            "Слов",
            "Предложений",
            "Символов без пробелов",
            "Символов с пробелами",
        ]
        for col, header in enumerate(headers, 1):
            cell = stats_sheet.cell(row=3, column=col, value=header)
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        if user_stats:
            # Основные данные
            row = 4
            type_name = ""
            if user_obj.role == "translator":
                type_name = "Назначенные предложения"
            elif user_obj.role == "corrector":
                type_name = "Проверенные предложения"
            else:
                type_name = "Загруженные документы"
            stats_sheet.cell(row=row, column=1, value=type_name)
            stats_sheet.cell(row=row, column=2, value=user_stats.get("total_words", 0))
            stats_sheet.cell(row=row, column=3, value=user_stats.get("total_sentences", 0))
            stats_sheet.cell(
                row=row,
                column=4,
                value=user_stats.get("total_characters_without_spaces", 0),
            )
            stats_sheet.cell(row=row, column=5, value=user_stats.get("total_characters", 0))
            # Данные по переводам (если есть)
            if user_stats.get("translated_sentences", 0) > 0:
                row += 1
                if user_obj.role == "translator":
                    type_name = "Выполненные переводы"
                elif user_obj.role == "corrector":
                    type_name = "Проверенные переводы"
                else:
                    type_name = "Переведенные предложения"
                stats_sheet.cell(row=row, column=1, value=type_name)
                stats_sheet.cell(row=row, column=2, value=user_stats.get("total_translated_words", 0))
                stats_sheet.cell(row=row, column=3, value=user_stats.get("translated_sentences", 0))
                stats_sheet.cell(
                    row=row,
                    column=4,
                    value=user_stats.get("total_translated_characters_without_spaces", 0),
                )
                stats_sheet.cell(
                    row=row,
                    column=5,
                    value=user_stats.get("total_translated_characters", 0),
                )
        # Лист с последними активностями
        activity_sheet = workbook.create_sheet("Последние активности")
        # Заголовок
        activity_sheet.merge_cells("A1:E1")
        activity_title_cell = activity_sheet["A1"]
        activity_title_cell.value = "Последние активности"
        activity_title_cell.font = openpyxl.styles.Font(size=16, bold=True)
        activity_title_cell.alignment = openpyxl.styles.Alignment(horizontal="center")
        if user_obj.role in ["admin", "representative"] and context_data.get("user_documents"):
            # Последние документы
            activity_sheet["A3"] = "Последние документы"
            activity_sheet["A3"].font = openpyxl.styles.Font(size=14, bold=True)
            doc_headers = ["Название", "Дата загрузки"]
            for col, header in enumerate(doc_headers, 1):
                cell = activity_sheet.cell(row=5, column=col, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            for row, document in enumerate(context_data["user_documents"][:5], 6):
                activity_sheet.cell(row=row, column=1, value=document.title)
                activity_sheet.cell(
                    row=row,
                    column=2,
                    value=document.uploaded_at.strftime("%d.%m.%Y %H:%M"),
                )
        elif user_obj.role == "translator":
            # Последние переводы
            if context_data.get("user_translations"):
                activity_sheet["A3"] = "Последние переводы"
                activity_sheet["A3"].font = openpyxl.styles.Font(size=14, bold=True)
                trans_headers = ["Документ", "Оригинал", "Перевод", "Статус", "Дата"]
                for col, header in enumerate(trans_headers, 1):
                    cell = activity_sheet.cell(row=5, column=col, value=header)
                    cell.font = openpyxl.styles.Font(bold=True)
                    cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                for row, translation in enumerate(context_data["user_translations"][:5], 6):
                    activity_sheet.cell(row=row, column=1, value=translation.sentence.document.title)
                    activity_sheet.cell(
                        row=row,
                        column=2,
                        value=(
                            translation.sentence.original_text[:50] + "..."
                            if len(translation.sentence.original_text) > 50
                            else translation.sentence.original_text
                        ),
                    )
                    activity_sheet.cell(
                        row=row,
                        column=3,
                        value=(
                            translation.translated_text[:50] + "..."
                            if len(translation.translated_text) > 50
                            else translation.translated_text
                        ),
                    )
                    status_map = {
                        "approved": "Утверждено",
                        "rejected": "Отклонено",
                        "pending": "На проверке",
                    }
                    activity_sheet.cell(
                        row=row,
                        column=4,
                        value=status_map.get(translation.status, translation.status),
                    )
                    activity_sheet.cell(
                        row=row,
                        column=5,
                        value=translation.translated_at.strftime("%d.%m.%Y %H:%M"),
                    )
        elif user_obj.role == "corrector":
            # Последние проверки
            if context_data.get("user_corrections"):
                activity_sheet["A3"] = "Последние проверки"
                activity_sheet["A3"].font = openpyxl.styles.Font(size=14, bold=True)
                corr_headers = ["Документ", "Оригинал", "Перевод", "Статус", "Дата"]
                for col, header in enumerate(corr_headers, 1):
                    cell = activity_sheet.cell(row=5, column=col, value=header)
                    cell.font = openpyxl.styles.Font(bold=True)
                    cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                for row, correction in enumerate(context_data["user_corrections"][:5], 6):
                    activity_sheet.cell(row=row, column=1, value=correction.sentence.document.title)
                    activity_sheet.cell(
                        row=row,
                        column=2,
                        value=(
                            correction.sentence.original_text[:50] + "..."
                            if len(correction.sentence.original_text) > 50
                            else correction.sentence.original_text
                        ),
                    )
                    activity_sheet.cell(
                        row=row,
                        column=3,
                        value=(
                            correction.translated_text[:50] + "..."
                            if len(correction.translated_text) > 50
                            else correction.translated_text
                        ),
                    )
                    status_map = {
                        "approved": "Утверждено",
                        "rejected": "Отклонено",
                        "pending": "На проверке",
                    }
                    activity_sheet.cell(
                        row=row,
                        column=4,
                        value=status_map.get(correction.status, correction.status),
                    )
                    activity_sheet.cell(
                        row=row,
                        column=5,
                        value=(
                            correction.corrected_at.strftime("%d.%m.%Y %H:%M")
                            if correction.corrected_at
                            else correction.translated_at.strftime("%d.%m.%Y %H:%M")
                        ),
                    )
        # Автоматическая ширина столбцов для всех листов
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for column in worksheet.columns:
                max_length = 0
                # Проверяем, что ячейка не является объединенной
                if hasattr(column[0], "column_letter"):
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except Exception:
                            continue
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width

        # Лист с предложениями
        sentences_sheet = workbook.create_sheet("Предложения")

        # Заголовок
        sentences_sheet.merge_cells("A1:F1")
        sentences_title_cell = sentences_sheet["A1"]
        sentences_title_cell.value = "Предложения"
        sentences_title_cell.font = openpyxl.styles.Font(size=16, bold=True)
        sentences_title_cell.alignment = openpyxl.styles.Alignment(horizontal="center")

        # Заголовки таблицы
        if user_obj.role == "translator":
            headers = [
                "Документ",
                "№ предложения",
                "Оригинальный текст",
                "Перевод",
                "Статус",
                "Дата перевода",
            ]
        elif user_obj.role == "corrector":
            headers = [
                "Документ",
                "№ предложения",
                "Оригинальный текст",
                "Перевод",
                "Статус",
                "Дата проверки",
            ]
        else:
            headers = [
                "Документ",
                "№ предложения",
                "Оригинальный текст",
                "Перевод",
                "Статус",
                "Дата",
            ]

        for col, header in enumerate(headers, 1):
            cell = sentences_sheet.cell(row=3, column=col, value=header)
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        # Данные предложений
        row = 4
        if user_obj.role == "translator":
            # Предложения, назначенные переводчику
            user_sentences = context_data.get("user_sentences", [])
            for sentence in user_sentences:
                sentences_sheet.cell(row=row, column=1, value=sentence.document.title)
                sentences_sheet.cell(row=row, column=2, value=sentence.sentence_number)
                sentences_sheet.cell(row=row, column=3, value=sentence.original_text)

                # Перевод
                if hasattr(sentence, "translation") and sentence.translation:
                    sentences_sheet.cell(row=row, column=4, value=sentence.translation.translated_text)
                    status_map = {
                        "approved": "Утверждено",
                        "rejected": "Отклонено",
                        "pending": "На проверке",
                    }
                    sentences_sheet.cell(
                        row=row,
                        column=5,
                        value=status_map.get(sentence.translation.status, sentence.translation.status),
                    )
                    sentences_sheet.cell(
                        row=row,
                        column=6,
                        value=sentence.translation.translated_at.strftime("%d.%m.%Y %H:%M"),
                    )
                else:
                    sentences_sheet.cell(row=row, column=4, value="Не переведено")
                    sentences_sheet.cell(row=row, column=5, value="Не назначено")
                    sentences_sheet.cell(row=row, column=6, value="")
                row += 1

        elif user_obj.role == "corrector":
            # Предложения, проверенные корректором
            user_corrections = context_data.get("user_corrections", [])
            for correction in user_corrections:
                sentence = correction.sentence
                sentences_sheet.cell(row=row, column=1, value=sentence.document.title)
                sentences_sheet.cell(row=row, column=2, value=sentence.sentence_number)
                sentences_sheet.cell(row=row, column=3, value=sentence.original_text)
                sentences_sheet.cell(row=row, column=4, value=correction.translated_text)

                status_map = {
                    "approved": "Утверждено",
                    "rejected": "Отклонено",
                    "pending": "На проверке",
                }
                sentences_sheet.cell(
                    row=row,
                    column=5,
                    value=status_map.get(correction.status, correction.status),
                )
                sentences_sheet.cell(
                    row=row,
                    column=6,
                    value=(
                        correction.corrected_at.strftime("%d.%m.%Y %H:%M")
                        if correction.corrected_at
                        else correction.translated_at.strftime("%d.%m.%Y %H:%M")
                    ),
                )
                row += 1

        else:
            # Для admin и representative - все предложения из загруженных документов
            user_documents = context_data.get("user_documents", [])
            for document in user_documents:
                for sentence in document.sentences.all():
                    sentences_sheet.cell(row=row, column=1, value=document.title)
                    sentences_sheet.cell(row=row, column=2, value=sentence.sentence_number)
                    sentences_sheet.cell(row=row, column=3, value=sentence.original_text)

                    # Перевод
                    if hasattr(sentence, "translation") and sentence.translation:
                        sentences_sheet.cell(
                            row=row,
                            column=4,
                            value=sentence.translation.translated_text,
                        )
                        status_map = {
                            "approved": "Утверждено",
                            "rejected": "Отклонено",
                            "pending": "На проверке",
                        }
                        sentences_sheet.cell(
                            row=row,
                            column=5,
                            value=status_map.get(sentence.translation.status, sentence.translation.status),
                        )
                        sentences_sheet.cell(
                            row=row,
                            column=6,
                            value=sentence.translation.translated_at.strftime("%d.%m.%Y %H:%M"),
                        )
                    else:
                        sentences_sheet.cell(row=row, column=4, value="Не переведено")
                        sentences_sheet.cell(row=row, column=5, value="Не назначено")
                        sentences_sheet.cell(row=row, column=6, value="")
                    row += 1

        workbook.save(output_path)
        return output_path
    except Exception as e:
        logging.error(f"Ошибка в export_user_report_to_xlsx: {str(e)}")
        raise


def export_user_report(user_obj, user_stats, context_data, format_type: str = "xlsx") -> str:
    """
    Основная функция для экспорта отчета по пользователю
    """
    try:
        # Создаем папку для экспорта, если её нет
        export_dir = os.path.join(default_storage.location, "exports")
        os.makedirs(export_dir, exist_ok=True)

        # Генерируем имя файла
        base_name = f"user_report_{user_obj.id}_{user_obj.username.replace(' ', '_')}"

        if format_type == "xlsx":
            output_path = os.path.join(export_dir, f"{base_name}.xlsx")
            return export_user_report_to_xlsx(user_obj, user_stats, context_data, output_path)
        else:
            raise ValueError(f"Неподдерживаемый формат экспорта: {format_type}")
    except Exception as e:
        logging.error(f"Ошибка в export_user_report: {str(e)}")
        raise
