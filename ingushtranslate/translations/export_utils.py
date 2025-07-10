import os
from typing import Dict, List

import docx
import openpyxl
from django.core.files.storage import default_storage

from .models import Document, Sentence, Translation


def export_to_txt(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в TXT файл
    """
    sentences = document.sentences.all().order_by("sentence_number")

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(f"Перевод документа: {document.title}\n")
        file.write("=" * 50 + "\n\n")

        for sentence in sentences:
            file.write(f"Предложение {sentence.sentence_number}:\n")
            file.write(f"Оригинал: {sentence.original_text}\n")

            if sentence.has_translation:
                translation = sentence.translation
                file.write(f"Перевод: {translation.translated_text}\n")
                if translation.translator_notes:
                    file.write(f"Заметки переводчика: {translation.translator_notes}\n")
                if translation.corrector_notes:
                    file.write(f"Заметки корректора: {translation.corrector_notes}\n")
                file.write(
                    f"Статус: {'Одобрен' if translation.is_approved else 'Отклонен' if translation.is_rejected else 'На проверке'}\n"
                )
            else:
                file.write("Перевод: Не переведено\n")

            file.write("-" * 30 + "\n\n")

    return output_path


def export_to_docx(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в DOCX файл
    """
    doc = docx.Document()

    # Заголовок
    title = doc.add_heading(f"Перевод документа: {document.title}", 0)
    doc.add_paragraph("=" * 50)
    doc.add_paragraph()

    sentences = document.sentences.all().order_by("sentence_number")

    for sentence in sentences:
        # Номер предложения
        doc.add_heading(f"Предложение {sentence.sentence_number}", level=1)

        # Оригинал
        doc.add_paragraph("Оригинал:", style="Heading 2")
        doc.add_paragraph(sentence.original_text)

        # Перевод
        doc.add_paragraph("Перевод:", style="Heading 2")
        if sentence.has_translation:
            translation = sentence.translation
            doc.add_paragraph(translation.translated_text)

            # Заметки
            if translation.translator_notes:
                doc.add_paragraph("Заметки переводчика:", style="Heading 3")
                doc.add_paragraph(translation.translator_notes)

            if translation.corrector_notes:
                doc.add_paragraph("Заметки корректора:", style="Heading 3")
                doc.add_paragraph(translation.corrector_notes)

            # Статус
            status_text = (
                "Одобрен"
                if translation.is_approved
                else "Отклонен" if translation.is_rejected else "На проверке"
            )
            doc.add_paragraph(f"Статус: {status_text}")
        else:
            doc.add_paragraph("Не переведено")

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def export_to_xlsx(document: Document, output_path: str) -> str:
    """
    Экспортирует переводы документа в XLSX файл
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Переводы"

    # Заголовки
    headers = [
        "Номер предложения",
        "Оригинальный текст",
        "Переведенный текст",
        "Переводчик",
        "Корректор",
        "Статус",
        "Заметки переводчика",
        "Заметки корректора",
        "Дата перевода",
        "Дата корректировки",
    ]

    for col, header in enumerate(headers, 1):
        sheet.cell(row=1, column=col, value=header)

    # Данные
    sentences = document.sentences.all().order_by("sentence_number")

    for row, sentence in enumerate(sentences, 2):
        sheet.cell(row=row, column=1, value=sentence.sentence_number)
        sheet.cell(row=row, column=2, value=sentence.original_text)

        if sentence.has_translation:
            translation = sentence.translation
            sheet.cell(row=row, column=3, value=translation.translated_text)
            sheet.cell(row=row, column=4, value=str(translation.translator))
            sheet.cell(
                row=row,
                column=5,
                value=str(translation.corrector) if translation.corrector else "",
            )

            # Статус
            if translation.is_approved:
                status = "Одобрен"
            elif translation.is_rejected:
                status = "Отклонен"
            else:
                status = "На проверке"
            sheet.cell(row=row, column=6, value=status)

            sheet.cell(row=row, column=7, value=translation.translator_notes)
            sheet.cell(row=row, column=8, value=translation.corrector_notes)
            sheet.cell(
                row=row,
                column=9,
                value=translation.translated_at.strftime("%Y-%m-%d %H:%M"),
            )
            sheet.cell(
                row=row,
                column=10,
                value=(
                    translation.corrected_at.strftime("%Y-%m-%d %H:%M")
                    if translation.corrected_at
                    else ""
                ),
            )
        else:
            sheet.cell(row=row, column=3, value="Не переведено")

    # Автоматическая ширина столбцов
    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        sheet.column_dimensions[column_letter].width = adjusted_width

    workbook.save(output_path)
    return output_path


def export_document_translations(document: Document, format_type: str) -> str:
    """
    Основная функция для экспорта переводов документа
    """
    # Создаем папку для экспорта, если её нет
    export_dir = os.path.join(default_storage.location, "exports")
    os.makedirs(export_dir, exist_ok=True)

    # Генерируем имя файла
    base_name = f"translation_{document.id}_{document.title.replace(' ', '_')}"

    if format_type == "txt":
        output_path = os.path.join(export_dir, f"{base_name}.txt")
        return export_to_txt(document, output_path)
    elif format_type == "docx":
        output_path = os.path.join(export_dir, f"{base_name}.docx")
        return export_to_docx(document, output_path)
    elif format_type == "xlsx":
        output_path = os.path.join(export_dir, f"{base_name}.xlsx")
        return export_to_xlsx(document, output_path)
    else:
        raise ValueError(f"Неподдерживаемый формат экспорта: {format_type}")


def get_document_statistics(document: Document) -> Dict:
    """
    Возвращает статистику по документу
    """
    total_sentences = document.sentences.count()
    translated_sentences = document.sentences.filter(translation__isnull=False).count()
    approved_sentences = document.sentences.filter(
        translation__is_approved=True
    ).count()
    rejected_sentences = document.sentences.filter(
        translation__is_rejected=True
    ).count()
    pending_sentences = document.sentences.filter(
        translation__isnull=False,
        translation__is_approved=False,
        translation__is_rejected=False,
    ).count()

    return {
        "total": total_sentences,
        "translated": translated_sentences,
        "approved": approved_sentences,
        "rejected": rejected_sentences,
        "pending": pending_sentences,
        "completion_percentage": round(
            (
                (translated_sentences / total_sentences * 100)
                if total_sentences > 0
                else 0
            ),
            2,
        ),
        "approval_percentage": round(
            (approved_sentences / total_sentences * 100) if total_sentences > 0 else 0,
            2,
        ),
    }
